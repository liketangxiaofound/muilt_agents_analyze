"""
文件解析工具

提供多格式文件解析能力，输出统一的数据结构。
"""
import os
import json
import csv
import logging
from io import StringIO

from strands import tool

# 可选依赖
try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None


def _resolve_path(file_path: str) -> str:
    from app.tools.storage_tools import process_file_key
    return process_file_key(file_path)


@tool
def parse_excel(file_path: str, sheet_name: str = "") -> str:
    """解析 Excel 文件（xlsx/xls），大文件只返回前 50 行预览。"""
    if openpyxl is None:
        return json.dumps({"error": "openpyxl 未安装"}, ensure_ascii=False)
    max_preview = 50
    try:
        full_path = _resolve_path(file_path)
        wb = openpyxl.load_workbook(full_path, data_only=True)
        sheets_data = []
        target_sheets = [sheet_name] if sheet_name else [wb.sheetnames[0]]
        for sname in target_sheets:
            if sname not in wb.sheetnames:
                continue
            ws = wb[sname]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(h) if h is not None else f"col_{i}" for i, h in enumerate(rows[0])]
            all_rows = [[None if cell is None else str(cell) if isinstance(cell, str) else cell for cell in row] for row in rows[1:]]
            total = len(all_rows)
            preview = all_rows[:max_preview]
            sheets_data.append({"sheet_name": sname, "row_count": total, "column_count": len(headers),
                "columns": headers, "rows": preview,
                "truncated": total > max_preview,
                "message": f"共 {total} 行，仅返回前 {max_preview} 行预览。" if total > max_preview else ""})
        return json.dumps({"type": "tabular", "source": os.path.basename(file_path), "sheets": sheets_data}, ensure_ascii=False, default=str)
    except Exception as e:
        return json.dumps({"error": f"Excel 解析失败: {str(e)}"}, ensure_ascii=False)


@tool
def parse_csv(file_path: str, delimiter: str = ",", has_header: bool = True) -> str:
    """解析 CSV 文件。大文件只返回前 50 行预览，避免上下文溢出。"""
    max_preview = 50
    try:
        full_path = _resolve_path(file_path)
        with open(full_path, "r", encoding="utf-8-sig") as f:
            content = f.read()
        reader = csv.reader(StringIO(content), delimiter=delimiter)
        rows = list(reader)
        if not rows:
            return json.dumps({"error": "CSV 文件为空"}, ensure_ascii=False)
        headers = rows[0] if has_header else [f"col_{i}" for i in range(len(rows[0]))]
        data_rows = rows[1:] if has_header else rows
        total = len(data_rows)
        preview = data_rows[:max_preview]
        return json.dumps({"type": "tabular", "source": os.path.basename(file_path),
            "sheets": [{"sheet_name": "data", "row_count": total, "column_count": len(headers),
                "columns": headers, "rows": preview,
                "truncated": total > max_preview,
                "message": f"共 {total} 行，仅返回前 {max_preview} 行预览。统计分析工具会使用完整数据。" if total > max_preview else ""}]},
            ensure_ascii=False)
    except UnicodeDecodeError:
        return _parse_csv_gbk(full_path, delimiter, has_header, max_preview)
    except Exception as e:
        return json.dumps({"error": f"CSV 解析失败: {str(e)}"}, ensure_ascii=False)


def _parse_csv_gbk(full_path, delimiter, has_header, max_preview):
    try:
        with open(full_path, "r", encoding="gbk") as f:
            content = f.read()
        reader = csv.reader(StringIO(content), delimiter=delimiter)
        rows = list(reader)
        if not rows:
            return json.dumps({"error": "CSV 文件为空"}, ensure_ascii=False)
        headers = rows[0] if has_header else [f"col_{i}" for i in range(len(rows[0]))]
        data_rows = rows[1:] if has_header else rows
        total = len(data_rows)
        preview = data_rows[:max_preview]
        return json.dumps({"type": "tabular", "source": os.path.basename(full_path),
            "sheets": [{"sheet_name": "data", "row_count": total, "column_count": len(headers),
                "columns": headers, "rows": preview,
                "truncated": total > max_preview,
                "message": f"共 {total} 行，仅返回前 {max_preview} 行预览。" if total > max_preview else ""}]},
            ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"CSV 解析失败(GBK): {str(e)}"}, ensure_ascii=False)


@tool
def parse_docx(file_path: str) -> str:
    """解析 Word 文档（docx），提取段落和表格。"""
    if docx is None:
        return json.dumps({"error": "python-docx 未安装"}, ensure_ascii=False)
    try:
        full_path = _resolve_path(file_path)
        document = docx.Document(full_path)
        paragraphs = [{"text": p.text.strip(), "style": p.style.name if p.style else "Normal"} for p in document.paragraphs if p.text.strip()]
        tables = []
        for i, table in enumerate(document.tables):
            rows_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            headers = rows_data[0] if rows_data else []
            tables.append({"table_index": i, "row_count": len(rows_data) - 1, "columns": headers, "rows": rows_data[1:]})
        return json.dumps({"type": "textual", "source": os.path.basename(file_path), "paragraph_count": len(paragraphs), "table_count": len(tables), "paragraphs": paragraphs, "tables": tables}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"Docx 解析失败: {str(e)}"}, ensure_ascii=False)


@tool
def parse_pdf(file_path: str) -> str:
    """解析 PDF 文件，提取文本内容。"""
    if pypdf is None:
        return json.dumps({"error": "pypdf 未安装"}, ensure_ascii=False)
    try:
        full_path = _resolve_path(file_path)
        reader = pypdf.PdfReader(full_path)
        pages, full_text = [], []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text = text.strip()
                pages.append({"page": i + 1, "text": text})
                full_text.append(text)
        return json.dumps({"type": "textual", "source": os.path.basename(file_path), "page_count": len(reader.pages), "pages": pages, "full_text": "\n\n".join(full_text)}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"PDF 解析失败: {str(e)}"}, ensure_ascii=False)


@tool
def detect_structure(parsed_data_json: str) -> str:
    """分析已解析数据的结构：识别列类型（数值/日期/分类/文本）。"""
    from datetime import datetime as dt
    try:
        data = json.loads(parsed_data_json)
    except json.JSONDecodeError:
        return json.dumps({"error": "输入不是有效 JSON"}, ensure_ascii=False)

    structure = {"type": data.get("type", "unknown"), "tables_analysis": [], "text_analysis": {}}
    for sheet in data.get("sheets", []):
        columns = sheet.get("columns", [])
        rows = sheet.get("rows", [])
        if not columns or not rows:
            continue
        col_types = {}
        for i, col in enumerate(columns):
            values = [row[i] for row in rows if i < len(row) and row[i] is not None]
            if not values:
                col_types[col] = "empty"
                continue
            numeric_count = sum(1 for v in values if isinstance(v, (int, float)))
            date_count = 0
            for v in values:
                if isinstance(v, str):
                    for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%d/%m/%Y"]:
                        try:
                            dt.strptime(v, fmt)
                            date_count += 1
                            break
                        except ValueError:
                            pass
            total = len(values)
            if numeric_count / total > 0.8:
                col_types[col] = "numeric"
            elif date_count / total > 0.5:
                col_types[col] = "date"
            elif len(set(str(v)[:50] for v in values)) / total < 0.3:
                col_types[col] = "category"
            else:
                col_types[col] = "text"
        structure["tables_analysis"].append({
            "sheet_name": sheet.get("sheet_name", ""),
            "row_count": sheet.get("row_count", 0),
            "column_count": len(columns),
            "column_types": col_types,
            "numeric_columns": [c for c, t in col_types.items() if t == "numeric"],
            "date_columns": [c for c, t in col_types.items() if t == "date"],
            "category_columns": [c for c, t in col_types.items() if t == "category"],
        })
    para_count = data.get("paragraph_count", 0)
    page_count = data.get("page_count", 0)
    if para_count > 0 or page_count > 0:
        structure["text_analysis"] = {"paragraph_count": para_count, "page_count": page_count, "table_count": data.get("table_count", 0)}
    return json.dumps(structure, ensure_ascii=False)
