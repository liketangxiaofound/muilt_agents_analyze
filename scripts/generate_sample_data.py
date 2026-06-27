"""
样例数据生成脚本

产出固定格式的测试文件到 examples/ 目录。
所有数据都是"结果导向"的——趋势、异常、排名都可预期，
确保分析结果稳定可复现。

运行: python scripts/generate_sample_data.py
"""
import os
import sys
import csv
import random
from datetime import datetime, timedelta

# 确保项目根目录在 sys.path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

EXAMPLES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "examples")
os.makedirs(EXAMPLES_DIR, exist_ok=True)

# 固定随机种子，确保每次生成的数据一致
random.seed(42)

# ============================================================
# 1. 财务支出数据
# ============================================================
def generate_finance_data():
    """生成财务支出 Excel，12个月 × 8个科目，含3个异常值"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        print("  跳过: openpyxl 未安装")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "财务支出明细"

    # 表头
    categories = ["办公费", "差旅费", "培训费", "设备购置", "维修维护", "水电物业", "咨询费", "其他支出"]
    months = [f"2025年{m:02d}月" for m in range(1, 13)]

    ws.append(["月份"] + categories)

    # 生成数据：大部分正常，个别异常
    base_values = {
        "办公费": 12000, "差旅费": 15000, "培训费": 8000,
        "设备购置": 20000, "维修维护": 6000, "水电物业": 10000,
        "咨询费": 18000, "其他支出": 5000,
    }
    # 波动幅度
    variances = {
        "办公费": 2000, "差旅费": 4000, "培训费": 3000,
        "设备购置": 8000, "维修维护": 1500, "水电物业": 1000,
        "咨询费": 5000, "其他支出": 1500,
    }

    anomaly_months = {
        3: {"设备购置": 85000},   # 3月设备购置异常偏高
        7: {"差旅费": 52000},     # 7月差旅费异常偏高
        11: {"咨询费": 62000},    # 11月咨询费异常偏高
    }

    for mi, month in enumerate(months):
        row = [month]
        for cat in categories:
            base = base_values[cat]
            var = variances[cat]
            # 正常波动 ±var
            val = base + random.randint(-var, var)
            # 叠加异常值
            if (mi + 1) in anomaly_months and cat in anomaly_months[mi + 1]:
                val = anomaly_months[mi + 1][cat]
            # 整体轻微上升趋势（每月增长0.5%）
            val = int(val * (1 + mi * 0.005))
            row.append(val)
        ws.append(row)

    # 样式
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # 自动列宽
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 20)

    path = os.path.join(EXAMPLES_DIR, "财务支出_2025.xlsx")
    wb.save(path)
    print(f"  ✓ 财务支出_2025.xlsx ({len(months)}行 × {len(categories)}列, 3个异常值)")


# ============================================================
# 2. 部门绩效数据
# ============================================================
def generate_performance_data():
    """生成部门绩效 Excel，6个部门 × 4个KPI，目标vs实际"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        print("  跳过: openpyxl 未安装")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "2025上半年绩效"

    departments = ["综合管理部", "财务部", "技术研发部", "市场拓展部", "人力资源部", "后勤保障部"]
    kpis = ["任务完成率(%)", "预算执行率(%)", "服务质量评分", "创新改进项数"]
    targets = {
        "综合管理部": [95, 98, 90, 5],
        "财务部": [98, 100, 92, 3],
        "技术研发部": [90, 95, 88, 8],
        "市场拓展部": [85, 90, 85, 6],
        "人力资源部": [92, 96, 88, 4],
        "后勤保障部": [93, 95, 86, 3],
    }

    ws.append(["部门"] + [f"{k}(目标)" for k in kpis] + [f"{k}(实际)" for k in kpis])

    for dept in departments:
        row = [dept]
        for i, target_val in enumerate(targets[dept]):
            # 实际值在目标值附近 ±10% 浮动
            if kpis[i] == "预算执行率(%)":
                actual = target_val + random.randint(-8, 2)  # 一般不超预算
            elif kpis[i] == "创新改进项数":
                actual = target_val + random.randint(-2, 4)
            else:
                actual = target_val + random.randint(-10, 8)
            row.append(target_val)
            row.append(max(0, actual))
        ws.append(row)

    # 样式
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 22)

    path = os.path.join(EXAMPLES_DIR, "部门绩效_2025H1.xlsx")
    wb.save(path)
    print(f"  ✓ 部门绩效_2025H1.xlsx ({len(departments)}部门 × {len(kpis)*2}指标)")


# ============================================================
# 3. 销售对比数据 (两个CSV)
# ============================================================
def generate_sales_data():
    """生成两个季度的销售 CSV，用于对比分析"""
    products = ["产品A", "产品B", "产品C", "产品D", "产品E"]
    regions = ["华东", "华南", "华北", "西南", "西北"]

    def generate_quarter_data(quarter_name: str, seed_offset: float) -> list:
        """生成一个季度的销售数据"""
        rows = []
        rows.append(["产品", "区域", "销售数量", "销售金额(万元)", "利润率(%)", "日期"])
        for product in products:
            for region in regions:
                qty = random.randint(50, 500)
                # Q2 比 Q1 略有增长
                if seed_offset > 0:
                    qty = int(qty * random.uniform(1.05, 1.25))
                price = round(random.uniform(0.5, 5.0), 2)
                amount = round(qty * price, 2)
                profit_rate = round(random.uniform(8, 35), 2)
                date = f"2025-{random.choice(range((seed_offset*3+1), (seed_offset*3+4))):02d}-{random.randint(1,28):02d}"
                rows.append([product, region, qty, amount, profit_rate, date])
        return rows

    # Q1 数据
    random.seed(42)
    q1_rows = generate_quarter_data("Q1", 0)
    path_q1 = os.path.join(EXAMPLES_DIR, "销售数据_2025Q1.csv")
    with open(path_q1, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(q1_rows)
    print(f"  ✓ 销售数据_2025Q1.csv ({len(q1_rows)-1}行 × {len(q1_rows[0])}列)")

    # Q2 数据
    q2_rows = generate_quarter_data("Q2", 1)
    path_q2 = os.path.join(EXAMPLES_DIR, "销售数据_2025Q2.csv")
    with open(path_q2, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(q2_rows)
    print(f"  ✓ 销售数据_2025Q2.csv ({len(q2_rows)-1}行 × {len(q2_rows[0])}列)")


# ============================================================
# 4. 统计分析样本 CSV
# ============================================================
def generate_statistics_csv():
    """生成多列数值的统计分析 CSV"""
    rows = []
    rows.append(["序号", "温度(℃)", "湿度(%)", "气压(hPa)", "风速(m/s)", "PM2.5(μg/m³)"])

    random.seed(42)
    for i in range(1, 61):
        temp = round(random.normalvariate(25, 5), 1)       # 均值25，标准差5
        humidity = round(random.normalvariate(60, 12), 1)   # 均值60，标准差12
        pressure = round(random.normalvariate(1013, 8), 1) # 均值1013，标准差8
        wind = round(random.weibullvariate(3, 1.5), 1)     # 偏态分布
        pm25 = round(random.lognormvariate(3.5, 0.5), 1)   # 对数正态

        # 方案几个异常点
        if i == 15:
            temp = 42.5  # 异常高温
        if i == 35:
            pm25 = 350.0  # 异常高PM2.5
        if i == 48:
            wind = 15.2   # 异常大风

        rows.append([i, temp, humidity, pressure, wind, pm25])

    path = os.path.join(EXAMPLES_DIR, "统计分析_样本.csv")
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f"  ✓ 统计分析_样本.csv ({len(rows)-1}行 × {len(rows[0])}列, 3个异常点)")


# ============================================================
# 5. Word 文档样例
# ============================================================
def generate_docx_sample():
    """生成结构化 Word 文档，含段落和表格"""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  跳过: python-docx 未安装")
        return

    doc = docx.Document()

    # 标题
    title = doc.add_heading("2025年度信息化建设工作报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("编制单位：信息化管理处")
    doc.add_paragraph(f"编制日期：2025年12月31日")
    doc.add_paragraph("")

    # 一、工作概述
    doc.add_heading("一、工作概述", level=2)
    doc.add_paragraph(
        "2025年，信息化管理处紧紧围绕单位中心工作，以\"数字政务、智慧管理\"为目标，"
        "全面推进信息化建设各项工作。全年完成信息化项目12个，预算执行率96.8%，"
        "系统可用率99.7%，用户满意度92.5分，较去年提升3.2个百分点。"
    )

    # 二、主要工作
    doc.add_heading("二、主要工作完成情况", level=2)

    sections = [
        ("（一）基础设施建设", "全年新增服务器12台，存储容量扩容至200TB，核心网络带宽升级至万兆。完成数据中心机房改造，PUE值从1.8降至1.5，年节约电费约18万元。"),
        ("（二）应用系统建设", "上线OA协同办公系统V3.0、财务管理系统、人事管理系统，完成数据互通互联。移动端应用覆盖率达到85%，实现移动审批、移动考勤等功能。"),
        ("（三）网络安全保障", "全年完成4次安全巡检、2次渗透测试、1次应急演练。发现并修复漏洞127个，其中高危漏洞12个。全年未发生重大网络安全事件。"),
        ("（四）数据治理", "编制《数据资源目录》，梳理数据资源536项。建立数据质量管理机制，修复数据质量问题3200余条。完成历史数据迁移和归档，释放存储空间约5TB。"),
    ]

    for title_text, content in sections:
        doc.add_heading(title_text, level=3)
        doc.add_paragraph(content)

    # 三、数据表格
    doc.add_heading("三、2025年度信息化经费支出统计", level=2)

    table = doc.add_table(rows=5, cols=5, style="Table Grid")
    headers = ["季度", "硬件采购(万元)", "软件采购(万元)", "运维服务(万元)", "合计(万元)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ("Q1", 85.2, 42.1, 28.5, 155.8),
        ("Q2", 63.5, 55.8, 30.2, 149.5),
        ("Q3", 48.9, 72.3, 32.1, 153.3),
        ("Q4", 120.5, 95.6, 35.8, 251.9),
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = str(val)

    doc.add_paragraph("")

    # 四、问题与不足
    doc.add_heading("四、问题与不足", level=2)
    issues = [
        "部分老旧系统技术债务较重，升级改造难度大。",
        "跨部门数据共享机制不够完善，存在信息孤岛现象。",
        "信息安全人才储备不足，高端安全运维人员短缺。",
    ]
    for issue in issues:
        doc.add_paragraph(issue, style="List Bullet")

    # 五、下一步计划
    doc.add_heading("五、2026年工作计划", level=2)
    plans = [
        "启动\"智慧政务\"一期项目，建设统一数据共享交换平台。",
        "完成老旧系统升级改造，迁移至云计算平台。",
        "建设AI辅助决策系统，提升数据分析应用能力。",
        "加强信息安全建设，通过ISO 27001信息安全管理体系认证。",
    ]
    for plan in plans:
        doc.add_paragraph(plan, style="List Bullet")

    path = os.path.join(EXAMPLES_DIR, "年度统计报告.docx")
    doc.save(path)
    print(f"  ✓ 年度统计报告.docx (4章节 + 1表格 + 建议措施)")


# ============================================================
# main
# ============================================================
def main():
    print("开始生成样例数据...\n")

    generate_finance_data()
    generate_performance_data()
    generate_sales_data()
    generate_statistics_csv()
    generate_docx_sample()

    print(f"\n样例数据已生成到: {EXAMPLES_DIR}")
    print(f"共 {len(os.listdir(EXAMPLES_DIR))} 个文件")


if __name__ == "__main__":
    main()
